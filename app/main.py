import os, tempfile, aiofiles, shutil, uuid
from typing import List, Optional, Dict, Any
from pathlib import Path
from fastapi import FastAPI, WebSocket, WebSocketDisconnect, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from PyPDF2 import PdfReader
import ocrmypdf
from pyannote.audio import Pipeline
from openai import OpenAI
from pydantic import BaseModel
from fpdf import FPDF
from pydub import AudioSegment
from PIL import Image
import os
from docx import Document
from fastapi.responses import StreamingResponse
import io
import fitz 
from docx import Document
from fpdf import FPDF
import httpx, asyncio, tempfile
from sse_starlette.sse import EventSourceResponse
from bs4 import BeautifulSoup
import json
# ---------- App & CORS ----------
app = FastAPI()
REDACCION_CHATS: Dict[str, Dict[str, Any]] = {}
origins = ["http://localhost:5173", "http://127.0.0.1:5173"]
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,           # viene de app/core/config.py
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
class RedaccionChatStart(BaseModel):
    documento_inicial: str
    evidence_id: Optional[str] = None
# ---------- Config ----------
DEFAULT_TEMPLATES = {
    "Demanda": "app/templates/demanda_base.txt",
    "Poder": "app/templates/poder_base.txt",
    "Escrito": "app/templates/escrito_base.txt",
    "Testamento": "app/templates/testamento_base.txt",
    "DeclaraciÃ³n Jurada": "app/templates/declaracion_Jurada_base.txt",
}

# ğŸš¨ API KEY fija (tuya)
client = OpenAI(api_key=os.getenv("APIGPT"))

# HuggingFace diarizaciÃ³n (opcional)
diarization_pipeline = Pipeline.from_pretrained(
    "pyannote/speaker-diarization",
    use_auth_token=os.getenv("HF_TOKEN")
)

BASE_DIR = Path("data")
PDF_DIR = BASE_DIR / "pdfs"
PDF_DIR.mkdir(parents=True, exist_ok=True)

EVIDENCES: Dict[str, str] = {}
chats: Dict[str, Dict[str, Any]] = {}

# =============== Utilidades ===============
def extract_text_from_pdf(path: str, max_pages: Optional[int] = None) -> str:
    reader = PdfReader(path)
    pages = range(len(reader.pages)) if max_pages is None else range(min(max_pages, len(reader.pages)))
    return "\n".join([reader.pages[i].extract_text() or "" for i in pages])

async def build_evidence_text(
    evidence_files: List[UploadFile],
    ocr: bool = True,
    max_pages_per_pdf: Optional[int] = None,
    max_chars_total: int = 120_000,
) -> str:
    if not evidence_files:
        return ""
    tmpdir = tempfile.mkdtemp(prefix="evidence_")
    parts: List[str] = []
    try:
        for idx, uf in enumerate(evidence_files, start=1):
            raw_path = os.path.join(tmpdir, f"raw_{idx}.pdf")
            async with aiofiles.open(raw_path, "wb") as f:
                data = await uf.read()
                await f.write(data)
            text_path = raw_path
            if ocr:
                ocr_path = os.path.join(tmpdir, f"ocr_{idx}.pdf")
                try:
                    ocrmypdf.ocr(raw_path, ocr_path, force_ocr=True, output_type="pdf", skip_text=False)
                    text_path = ocr_path
                except Exception:
                    text_path = raw_path
            text = extract_text_from_pdf(text_path, max_pages=max_pages_per_pdf).strip()
            header = f"=== PRUEBA {idx}: {uf.filename or 'adjunto.pdf'} ==="
            parts.append(f"{header}\n{text}\n")
        full = "\n".join(parts)
        if len(full) > max_chars_total:
            full = full[:max_chars_total] + "\n[...truncado por longitud]"
        return full
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)

def generar_pdf(path: str, messages: list[dict]):
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("DejaVu", "", "C:/Windows/Fonts/DejaVuSans.ttf", uni=True)
    pdf.set_font("DejaVu", size=12)

    last_doc = None
    for m in reversed(messages):
        if m["role"] == "assistant" and m["content"].strip():
            last_doc = m["content"].strip()
            break
    pdf.multi_cell(0, 10, last_doc or "[Documento aÃºn sin contenido]")
    pdf.output(path)

# =============== Chat API ===============
chat_history = {}
SPRINGBOOT_URL = "http://localhost:8080/api/messages"
SYSTEM_PROMPT = (
    "Eres un asistente jurÃ­dico especializado en legislaciÃ³n uruguaya. "
    "Debes responder como un abogado experto en derecho uruguayo, citando leyes, decretos y artÃ­culos aplicables. "
    "Si un usuario menciona una ley, decreto o resoluciÃ³n por nÃºmero (por ejemplo, 'Ley 12234' o 'Decreto 456/2012'), "
    "debes devolver un JSON con la acciÃ³n de bÃºsqueda, sin explicar nada todavÃ­a. "
    "El formato exacto debe ser: "
    '{"action":"fetch_page","url":"https://www.impo.com.uy/bases/leyes/{numero_de_ley}","keyword":"{numero_de_ley}"} '
    "(reemplazando {numero_de_ley} por el nÃºmero real mencionado). "
    "Si el nÃºmero es un decreto, usÃ¡ la URL https://www.impo.com.uy/bases/decretos/{numero}. "
    "Si el usuario no menciona una norma especÃ­fica, responde normalmente en texto plano. "
    "Nunca digas que no tenÃ©s acceso en lÃ­nea: si se trata de una ley uruguaya, asumÃ­ que puedes generar la URL oficial."
)
async def persist_to_spring(chat_id, message, response, files):
    async with httpx.AsyncClient() as http_client:
        data = {
            "chat_id": chat_id,
            "message": message,
            "response": response,
        }

        # Armamos la lista de archivos correctamente
        files_data = []
        for path in files:
            files_data.append(("files", (path.name, open(path, "rb"), "application/octet-stream")))

        await http_client.post(SPRINGBOOT_URL, data=data, files=files_data)




def scrape_page(url: str) -> str:
    """Descarga y limpia el texto de una pÃ¡gina web."""
    try:
        res = requests.get(url, timeout=10)
        res.raise_for_status()
        soup = BeautifulSoup(res.text, "html.parser")
        text = soup.get_text(separator="\n")
        # Limitar longitud para evitar contextos enormes
        return text[:15000]
    except Exception as e:
        return f"Error al intentar acceder a {url}: {e}"


@app.post("/chat/stream")
async def chat_stream(
    chat_id: str = Form(...),
    message: str = Form(...),
    files: list[UploadFile] = File(default=None)
):
    tmpdir = tempfile.mkdtemp(prefix="chat_")
    saved_files = []
    for file in (files or []):
        path = Path(tmpdir) / file.filename
        with open(path, "wb") as f:
            f.write(await file.read())
        saved_files.append(path)

    chat_history.setdefault(chat_id, [{"role": "system", "content": SYSTEM_PROMPT}])
    chat_history[chat_id].append({"role": "user", "content": message})

    # Primera respuesta: ver si GPT pide una fuente externa
    response = client.chat.completions.create(
        model="gpt-5",
        messages=chat_history[chat_id]
    )
    content = response.choices[0].message.content.strip()

    # Intentamos interpretar si GPT devolviÃ³ JSON con acciÃ³n fetch_page
    try:
        parsed = json.loads(content)
        if parsed.get("action") == "fetch_page" and parsed.get("url"):
            url = parsed["url"]
            keyword = parsed.get("keyword", "")
            scraped_text = scrape_page(url)

            # Segunda consulta: GPT analiza el contenido descargado
            followup = client.chat.completions.create(
                model="gpt-5",
                messages=[
                    {"role": "system", "content": "Eres un abogado uruguayo, resume y explica el texto legal proporcionado."},
                    {"role": "user", "content": f"Palabra clave: {keyword}\n\nTexto extraÃ­do de {url}:\n{scraped_text}"}
                ]
            )
            full_response = followup.choices[0].message.content
        else:
            full_response = content
    except json.JSONDecodeError:
        # No era JSON, GPT respondiÃ³ directamente
        full_response = content

    # Guardamos en el historial y en SpringBoot
    chat_history[chat_id].append({"role": "assistant", "content": full_response})
    await persist_to_spring(chat_id, message, full_response, saved_files)

    return {"response": full_response}

# =============== TranscripciÃ³n + DiarizaciÃ³n ===============
@app.post("/transcribir_diarizado/")
async def transcribir_diarizado(audio: UploadFile = File(...)):
    # Guardar archivo temporal
    tmp_path = os.path.join(tempfile.gettempdir(), audio.filename)
    with open(tmp_path, "wb") as f:
        f.write(await audio.read())

    # Convertir SIEMPRE a WAV estÃ¡ndar
    converted_path = tmp_path + ".wav"
    try:
        sound = AudioSegment.from_file(tmp_path)  # autodetecta formato
        sound.export(converted_path, format="wav", codec="pcm_s16le")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error convirtiendo audio: {str(e)}")

    # 1. Transcribir con GPT-4o-transcribe
    with open(converted_path, "rb") as f:
        transcription = client.audio.transcriptions.create(
            model="gpt-4o-transcribe",
            file=f
        )

    raw_text = transcription.text.strip()

    # 2. DiarizaciÃ³n con GPT-5
#     prompt = f"""
# Eres un asistente jurÃ­dico que organiza transcripciones de audios judiciales.
# El siguiente texto es una transcripciÃ³n cruda.
# Asigna nombres/roles apropiados a cada hablante, si reconoces el nombre del participante por ejemplo : Pedro gonzalez unicamente coloca el nombre y no el rol (Juez, Fiscal, Testigo, Abogado, Acusado, etc).
# Formato esperado:

# Rol o nombre: texto
# Rol o nombre: texto
# ...

# Texto transcrito:
# {raw_text}
#     """

#     resp = client.chat.completions.create(
#         model="gpt-5",
#         messages=[{"role": "user", "content": prompt}],
#         max_completion_tokens=4096,
#     )

#     diarized_text = resp.choices[0].message.content.strip()

    return {"conversacion": raw_text}

# =============== WebSocket live chat ===============
@app.websocket("/ws/chat")
async def websocket_endpoint(websocket: WebSocket):
    await websocket.accept()
    history = [{"role": "system", "content": "Eres un asistente legal especializado en derecho uruguayo ."}]
    try:
        while True:
            data = await websocket.receive_text()
            history.append({"role": "user", "content": data})
            completion = client.chat.completions.create(
                model="gpt-5",   # âœ… ahora GPT-5
                messages=history,
                temperature=0.3,
            )
            reply = completion.choices[0].message.content
            history.append({"role": "assistant", "content": reply})
            await websocket.send_text(reply)
    except WebSocketDisconnect:
        print("Cliente desconectado")

@app.post("/ocr_archivo_con_texto/")
async def ocr_archivo_con_texto(archivo: UploadFile = File(...)):
    try:
        # Crear carpeta temporal
        tmpdir = tempfile.mkdtemp(prefix="ocr_")
        input_path = os.path.join(tmpdir, archivo.filename)

        # Guardar archivo subido
        with open(input_path, "wb") as f:
            f.write(await archivo.read())

        # Si es imagen -> convertir a PDF temporal
        ext = Path(archivo.filename).suffix.lower()
        if ext in [".jpg", ".jpeg", ".png"]:
            img = Image.open(input_path).convert("RGB")
            input_pdf = os.path.join(tmpdir, "input.pdf")
            img.save(input_pdf, "PDF")
            input_path = input_pdf

        # Salida final con OCR
        output_pdf = os.path.join(tmpdir, "ocr_output.pdf")

        try:
            ocrmypdf.ocr(input_path, output_pdf, force_ocr=True, output_type="pdf")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error en OCR: {str(e)}")

        # Devolver archivo procesado
        return FileResponse(
            output_pdf,
            media_type="application/pdf",
            filename="archivo_ocr.pdf"
        )

    finally:
        # âš ï¸ Importante: no borramos tmpdir inmediatamente porque se necesita leer el PDF en FileResponse
        # FastAPI eliminarÃ¡ el archivo cuando termine la request.
        pass
@app.post("/resumir_documentos/")
async def resumir_documentos(
    files: List[UploadFile] = File(..., description="Lista de documentos a resumir"),
    ocr: bool = Form(True)
):
    informes = []

    tmpdir = tempfile.mkdtemp(prefix="resumenes_")
    try:
        for idx, uf in enumerate(files, start=1):
            # Guardar archivo temporal
            raw_path = os.path.join(tmpdir, uf.filename)
            async with aiofiles.open(raw_path, "wb") as f:
                data = await uf.read()
                await f.write(data)

            # Extraer texto (PDF o texto plano)
            if uf.filename.lower().endswith(".pdf"):
                text = extract_text_from_pdf(raw_path)
            else:
                try:
                    text = Path(raw_path).read_text(encoding="utf-8", errors="ignore")
                except:
                    text = ""

            # Resumir con GPT
            prompt = f"""
            ResumÃ­ lo mÃ¡s detalladamente posible el siguiente documento en un pÃ¡rrafo claro y conciso.
            Documento: {uf.filename}

            Contenido:
            {text[:100000]}  # limitar por seguridad
            """
            resp = client.chat.completions.create(
                model="gpt-5-mini",
                messages=[{"role": "user", "content": prompt}],
                max_completion_tokens=6000
            )
            resumen = resp.choices[0].message.content.strip()

            informes.append({
                "nombre": uf.filename,
                "resumen": resumen
            })

        informe_final = "\n\n".join([f"Documento: {inf['nombre']}\nResumen: {inf['resumen']}" for inf in informes])
        return {"documentos": informes, "informe_final": informe_final}

    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)
    
@app.post("/convertir_documento/")
async def convertir_documento(
    file: UploadFile = File(...),
    formato_salida: str = Form(..., regex="^(pdf|txt|docx)$")
):
    tmpdir = tempfile.mkdtemp(prefix="convert_")
    try:
        input_path = Path(tmpdir) / file.filename
        async with aiofiles.open(input_path, "wb") as f:
            data = await file.read()
            await f.write(data)

        output_path = Path(tmpdir) / f"convertido.{formato_salida}"

        ext_origen = input_path.suffix.lower()

        # PDF â†’ TXT
        if ext_origen == ".pdf" and formato_salida == "txt":
            text = ""
            with fitz.open(input_path) as pdf:
                for page in pdf:
                    text += page.get_text()
            output_path.write_text(text, encoding="utf-8")

        # DOCX â†’ TXT
        elif ext_origen == ".docx" and formato_salida == "txt":
            doc = Document(input_path)
            text = "\n".join(p.text for p in doc.paragraphs)
            output_path.write_text(text, encoding="utf-8")

        # TXT â†’ PDF
        elif ext_origen == ".txt" and formato_salida == "pdf":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            with open(input_path, "r", encoding="utf-8") as txt:
                for line in txt:
                    pdf.cell(200, 10, txt=line.encode("latin-1", "replace").decode("latin-1"), ln=True)
            pdf.output(str(output_path))

        # TXT â†’ DOCX
        elif ext_origen == ".txt" and formato_salida == "docx":
            text = input_path.read_text(encoding="utf-8")
            doc = Document()
            for line in text.splitlines():
                doc.add_paragraph(line)
            doc.save(str(output_path))

        # Si el formato de entrada y salida son iguales o no hay conversiÃ³n
        else:
            shutil.copy(str(input_path), str(output_path))

        return FileResponse(
            output_path,
            media_type="application/octet-stream",
            filename=f"convertido.{formato_salida}"
        )

    finally:
        pass
# ==================== CHAT DE REDACCIÃ“N ====================

PDF_DIR = Path("pdfs")
PDF_DIR.mkdir(exist_ok=True)
EVIDENCES: Dict[str, str] = {}
REDACCION_CHATS: Dict[str, Dict[str, Any]] = {}

@app.post("/redaccion_documento/")
async def redaccion_documento(
    document_type: str = Form(...),
    instructions: str = Form(""),
    template_file: Optional[UploadFile] = File(None),
    evidence_files: Optional[List[UploadFile]] = File(None),
):
    """Genera el primer borrador jurÃ­dico completo."""
    tmpdir = tempfile.mkdtemp(prefix="redaccion_")
    try:
        base_text = ""
        if template_file:
            path = os.path.join(tmpdir, template_file.filename)
            async with aiofiles.open(path, "wb") as f:
                await f.write(await template_file.read())
            base_text = Path(path).read_text(encoding="utf-8", errors="ignore").strip()

        evidence_text = ""
        if evidence_files:
            for ev in evidence_files:
                evidence_text += f"\n- {ev.filename}"

        prompt = f"""
RedactÃ¡ un documento jurÃ­dico completo de tipo **{document_type}**, siguiendo lenguaje jurÃ­dico uruguayo formal.
TenÃ© en cuenta, en ningÃºn caso podÃ©s inventar hechos o datos no provistos y si se quiere hacer referencia a una evidencia esta debe ser adjuntada directamente en el documento, es decir nunca hacer referencia a un pdf externo.
â€¢ Instrucciones: {instructions or '[sin instrucciones]'}
â€¢ Evidencias: {evidence_text or '[sin evidencias]'}
â€¢ Plantilla base: {base_text or '[sin plantilla]'}
El resultado debe ser un escrito jurÃ­dico formal y coherente, sin comentarios externos.
"""

        completion = client.chat.completions.create(
            model="gpt-5",
            messages=[
                {"role": "system", "content": "Sos un abogado uruguayo experto en derecho civil y procesal."},
                {"role": "user", "content": prompt},
            ],
            max_completion_tokens=10096,
        )

        generated_text = completion.choices[0].message.content.strip()

        evidence_id = str(uuid.uuid4())
        pdf_path = PDF_DIR / f"redaccion_{evidence_id}.pdf"

        # Genera PDF con el texto inicial
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, generated_text.encode("latin-1", "replace").decode("latin-1"))
        pdf.output(str(pdf_path))

        EVIDENCES[evidence_id] = evidence_text

        return {
            "generated_text": completion,
            "evidence_id": evidence_id,
            "pdf_url": f"http://127.0.0.1:8000/redaccion_chat/pdf/{evidence_id}",
        }
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


@app.post("/redaccion_chat/start/")
async def redaccion_chat_start(req: dict):
    """Inicia un chat con el documento ya generado."""
    chat_id = str(uuid.uuid4())
    REDACCION_CHATS[chat_id] = {
        "texto_actual": req["documento_inicial"],
        "pdf_path": str(PDF_DIR / f"redaccion_{req.get('evidence_id', chat_id)}.pdf"),
        "evidencia_text": EVIDENCES.get(req.get("evidence_id", ""), ""),
    }

    return {
        "chat_id": chat_id,
        "respuesta": "âœ… Documento generado correctamente. Ya podÃ©s empezar a editarlo.",
        "pdf_url": f"http://127.0.0.1:8000/redaccion_chat/pdf/{req.get('evidence_id', chat_id)}",
    }


@app.post("/redaccion_chat/message/")
async def redaccion_chat_message(req: dict):
    """
    Recibe instrucciones de modificaciÃ³n del usuario, actualiza el documento,
    y devuelve el feedback textual del chatbot + el nuevo PDF.
    """
    chat_id = req["chat_id"]
    mensaje = req["mensaje_usuario"]

    if chat_id not in REDACCION_CHATS:
        raise HTTPException(404, "Chat no encontrado")

    chat = REDACCION_CHATS[chat_id]
    texto_actual = chat["texto_actual"]

    # ğŸ§  Instrucciones para GPT
    prompt = f"""
El siguiente texto es un documento jurÃ­dico uruguayo en proceso de redacciÃ³n.
DebÃ©s aplicar la instrucciÃ³n del usuario al documento, manteniendo su tono formal y estructura.
DespuÃ©s de modificar el documento, explicÃ¡ brevemente (en 2-3 oraciones) quÃ© cambios realizaste.

Primero devolvÃ© el nuevo documento completo bajo la etiqueta [DOCUMENTO].
Luego devolvÃ© un resumen breve o confirmaciÃ³n bajo la etiqueta [CHATBOT].

Documento actual:
---
{texto_actual}
---

InstrucciÃ³n del usuario:
---
{mensaje}
---

Evidencias relevantes:
---
{chat.get('evidencia_text', '[sin evidencias]')}
---
"""

    completion = client.chat.completions.create(
        model="gpt-5",
        messages=[
            {"role": "system", "content": "Sos un abogado uruguayo experto en redacciÃ³n de documentos jurÃ­dicos."},
            {"role": "user", "content": prompt},
        ],
        max_completion_tokens=4096,
    )
   
    full_response = completion.choices[0].message.content.strip()
    print (full_response)
    # ğŸ§© Separar el documento del feedback
    new_text = ""
    feedback = "âœ… Documento actualizado."
    if "[DOCUMENTO]" in full_response:
        parts = full_response.split("[CHATBOT]")
        new_text = parts[0].replace("[DOCUMENTO]", "").strip()
        if len(parts) > 1:
            feedback = parts[1].strip()
            chat["texto_actual"] = new_text

    # ğŸ“„ Regenerar PDF actualizado
            pdf_path = Path(chat["pdf_path"])
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, new_text.encode("latin-1", "replace").decode("latin-1"))
            pdf.output(str(pdf_path))

    else:
        # fallback si el modelo no respetÃ³ las etiquetas
        new_text = full_response
        feedback = "OcurriÃ³ un error al interpretar la respuesta del modelo."

    # ğŸ–‹ï¸ Actualizar texto del documento en memoria

    # ğŸ” Responder al frontend
    return {
        "status": "âœ… OK",
        "message": "Asistente JurÃ­dico IA funcionando correctamente",
        "docs": "Visita /docs para ver la documentaciÃ³n interactiva",
    }
