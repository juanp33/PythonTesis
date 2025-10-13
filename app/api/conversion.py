from fastapi import APIRouter, UploadFile, File, Form
from fastapi.responses import FileResponse
from docx import Document
from fpdf import FPDF
import fitz, tempfile, aiofiles, shutil
from pathlib import Path

router = APIRouter(prefix="/convertir", tags=["Conversión"])

@router.post("/")
async def convertir_documento(file: UploadFile = File(...), formato_salida: str = Form(...)):
    tmpdir = tempfile.mkdtemp(prefix="conv_")
    try:
        input_path = Path(tmpdir) / file.filename
        async with aiofiles.open(input_path, "wb") as f:
            await f.write(await file.read())
        output_path = Path(tmpdir) / f"convertido.{formato_salida}"
        ext = input_path.suffix.lower()

        if ext == ".pdf" and formato_salida == "txt":
            text = "".join(page.get_text() for page in fitz.open(input_path))
            output_path.write_text(text, encoding="utf-8")

        elif ext == ".docx" and formato_salida == "txt":
            text = "\n".join(p.text for p in Document(input_path).paragraphs)
            output_path.write_text(text, encoding="utf-8")

        elif ext == ".txt" and formato_salida == "pdf":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            for line in input_path.read_text(encoding="utf-8").splitlines():
                pdf.cell(200, 10, txt=line.encode("latin-1", "replace").decode("latin-1"), ln=True)
            pdf.output(str(output_path))

        elif ext == ".txt" and formato_salida == "docx":
            text = input_path.read_text(encoding="utf-8")
            doc = Document()
            for line in text.splitlines():
                doc.add_paragraph(line)
            doc.save(str(output_path))

        else:
            shutil.copy(str(input_path), str(output_path))

        return FileResponse(output_path, filename=output_path.name)
    finally:
        pass
